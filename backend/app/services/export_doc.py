"""회의록 문서(.docx) 내보내기 — resource/doc/[회의록] 양식.hwp 레이아웃 재현.

.hwp(바이너리)는 프로그램으로 내용을 채울 수 없어, 같은 레이아웃의 .docx를 생성한다
(한글(HWP)·Word 모두에서 열림):

  회의록                                         ← 큰 제목
  ┌─────────┬───────────────┬────────┬────────────────┐
  │ 회의일시 │ 2026년 4월 …  │ 회의명 │ <제목> (#태그)  │
  ├─────────┼───────────────┴────────┴────────────────┤
  │ 참석자   │ 이름(소속 · 직책) 콤마 목록               │
  └─────────┴──────────────────────────────────────────┘
  ┌─────────┬──────────────────────────────────────────┐
  │ 회의내용 │ 1. 내용 / 핵심 내용 / 결정 사항 / 타임라인… │
  └─────────┴──────────────────────────────────────────┘
  ┌─────────┬──────────────────────────────────────────┐
  │ 특이사항 │ (빈칸 — 출력 후 직접 작성)                │
  └─────────┴──────────────────────────────────────────┘
"""

import json
import re
from datetime import datetime
from io import BytesIO

FONT = "맑은 고딕"
LABEL_BG = "E7E6E6"  # 라벨 셀 연회색
WEEKDAYS = ["월요일", "화요일", "수요일", "목요일", "금요일", "토요일", "일요일"]

_BOLD_MD_RE = re.compile(r"\*\*(.+?)\*\*")


def _fmt_datetime(iso: str | None) -> str:
    if not iso:
        return "-"
    try:
        d = datetime.fromisoformat(iso)
    except ValueError:
        return iso
    return f"{d.year}년 {d.month:02d}월 {d.day:02d}일 {WEEKDAYS[d.weekday()]} {d.hour:02d}:{d.minute:02d}"


def _fmt_clock(sec: float | None) -> str:
    s = int(sec or 0)
    return f"{s // 3600:02d}:{s % 3600 // 60:02d}:{s % 60:02d}"


def format_participants_grouped(participants: list[dict]) -> str:
    """참석자를 소속별로 묶어 '마인즈에이아이(박대한, 최영준), TTA(김인영)' 형식으로.

    소속이 없는 참석자는 마지막에 이름만 콤마로 덧붙인다. docx/pdf 공용.
    """
    grouped: dict[str, list[str]] = {}
    loose: list[str] = []
    for p in participants:
        name = str(p.get("name") or "").strip()
        if not name:
            continue
        org = str(p.get("organization") or "").strip()
        if org:
            grouped.setdefault(org, []).append(name)
        else:
            loose.append(name)
    parts = [f"{org}({', '.join(names)})" for org, names in grouped.items()]
    parts += loose
    return ", ".join(parts)


def _set_run(run, size=10, bold=False):
    from docx.shared import Pt

    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    # 한글 글꼴은 eastAsia에도 지정해야 적용된다
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(_qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(_qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(_qn("w:eastAsia"), FONT)


def _qn(tag):
    from docx.oxml.ns import qn

    return qn(tag)


def _shade_cell(cell, color: str):
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, text: str, size=10, bold=False, center=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(text), size=size, bold=bold)


def _label_cell(cell, text: str):
    from docx.enum.table import WD_ALIGN_VERTICAL

    _cell_text(cell, text, size=10.5, bold=True, center=True)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade_cell(cell, LABEL_BG)


def _add_line(cell, text: str, size=10, bold=False, indent_cm=0.0, before=0, after=2):
    """셀에 문단 한 줄 추가 — **굵게** 인라인 마크다운 지원."""
    from docx.shared import Cm, Pt

    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent_cm:
        pf.left_indent = Cm(indent_cm)
    pos = 0
    for m in _BOLD_MD_RE.finditer(text):
        if m.start() > pos:
            _set_run(p.add_run(text[pos : m.start()]), size=size, bold=bold)
        _set_run(p.add_run(m.group(1)), size=size, bold=True)
        pos = m.end()
    if pos < len(text):
        _set_run(p.add_run(text[pos:]), size=size, bold=bold)
    return p


def _add_markdown(cell, md: str):
    """discussion 마크다운(### 소제목 / - 불릿 / 문단)을 셀 문단으로 렌더."""
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            _add_line(cell, line[4:], size=10, bold=True, before=6)
        elif line.startswith("## "):
            _add_line(cell, line[3:], size=10.5, bold=True, before=8)
        elif line.startswith("- [x] ") or line.startswith("- [X] "):
            _add_line(cell, f"☑ {line[6:]}", indent_cm=0.25)
        elif line.startswith("- [ ] "):
            _add_line(cell, f"☐ {line[6:]}", indent_cm=0.25)
        elif line.startswith("- "):
            _add_line(cell, f"• {line[2:]}", indent_cm=0.25)
        else:
            _add_line(cell, line)


def _section(cell, no: int, title: str) -> None:
    _add_line(cell, f"{no}. {title}", size=11, bold=True, before=10, after=4)


def build_minutes_docx(
    meeting: dict,
    participants: list[dict],
    bookmarks: list[dict],
    summary: dict | None,
) -> bytes:
    from docx import Document
    from docx.shared import Cm, Pt

    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    usable_cm = 21.0 - 1.8 * 2  # A4 기준 본문 폭

    # 제목
    title_p = doc.add_paragraph()
    _set_run(title_p.add_run("회의록"), size=22, bold=True)
    title_p.paragraph_format.space_after = Pt(10)

    # ---- 표 1: 회의일시/회의명 + 참석기관 ----
    t1 = doc.add_table(rows=2, cols=4)
    t1.style = "Table Grid"
    widths = [Cm(2.6), Cm(6.4), Cm(2.2), Cm(usable_cm - 2.6 - 6.4 - 2.2)]
    for row in t1.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    _label_cell(t1.cell(0, 0), "회의일시")
    _cell_text(t1.cell(0, 1), _fmt_datetime(meeting.get("started_at")))
    _label_cell(t1.cell(0, 2), "회의명")
    tag = str(meeting.get("tag") or "").strip()
    meeting_name = str(meeting.get("title") or "")
    if tag:
        meeting_name += f"  (#{tag})"
    _cell_text(t1.cell(0, 3), meeting_name)

    _label_cell(t1.cell(1, 0), "참석자")
    body = t1.cell(1, 1).merge(t1.cell(1, 2)).merge(t1.cell(1, 3))
    _cell_text(body, format_participants_grouped(participants) or "-")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- 표 2: 회의내용 ----
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    t2.cell(0, 0).width = Cm(2.6)
    t2.cell(0, 1).width = Cm(usable_cm - 2.6)
    _label_cell(t2.cell(0, 0), "회의내용")

    content = t2.cell(0, 1)
    content.text = ""

    summary = summary or {}
    section_no = 1

    discussion = str(summary.get("discussion") or "").strip()
    if discussion:
        _section(content, section_no, "내용")
        section_no += 1
        _add_markdown(content, discussion)

    key_points = summary.get("key_points") or []
    if key_points:
        _section(content, section_no, "핵심 내용")
        section_no += 1
        for item in key_points:
            _add_line(content, f"• {item}", indent_cm=0.25)

    _section(content, section_no, "결정 사항")
    section_no += 1
    decisions = summary.get("decisions") or []
    if decisions:
        for item in decisions:
            _add_line(content, f"☑ {item}", indent_cm=0.25)
    else:
        _add_line(content, "명확히 확정된 결정사항은 없음", indent_cm=0.25)

    followups = summary.get("followups") or []
    if followups:
        _section(content, section_no, "추가 확인 필요 사항")
        section_no += 1
        for item in followups:
            _add_line(content, f"☐ {item}", indent_cm=0.25)

    timed = [b for b in bookmarks if b.get("kind") != "note"]
    if timed:
        _section(content, section_no, "타임라인")
        section_no += 1
        for b in sorted(timed, key=lambda x: float(x.get("time_sec") or 0)):
            _add_line(
                content,
                f"**{_fmt_clock(b.get('time_sec'))}** — {str(b.get('title') or '').strip()}",
                indent_cm=0.25,
            )

    if not discussion and not key_points and not decisions:
        _add_line(content, "요약이 아직 생성되지 않았습니다.")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- 표 3: 특이사항 (일반 메모) ----
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    t3.cell(0, 0).width = Cm(2.6)
    t3.cell(0, 1).width = Cm(usable_cm - 2.6)
    _label_cell(t3.cell(0, 0), "특이사항")
    # 특이사항은 비워 둔다 — 출력 후 직접 작성 (사용자 요청)
    _cell_text(t3.cell(0, 1), "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def parse_summary_row(row) -> dict | None:
    """summaries sqlite Row → build_minutes_docx용 dict (라우터에서 재사용)."""
    if row is None:
        return None
    return {
        "discussion": row["discussion"] or "",
        "key_points": json.loads(row["key_points"]),
        "decisions": json.loads(row["decisions"]),
        "followups": json.loads(row["followups"]) if row["followups"] else [],
    }
